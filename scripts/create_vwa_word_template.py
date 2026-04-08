#!/usr/bin/env python3
"""
Generate a professional VWA (Vorwissenschaftliche Arbeit) Word template
for BG Zehnergasse, matching the existing LaTeX template design.

Usage: python3 scripts/create_vwa_word_template.py
Output: public/templates/bgz-vwa-word.docx
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# Constants
# ============================================================
TEAL = RGBColor(0x00, 0x79, 0x6B)       # #00796b - Primary heading color
DARK_BLUE = RGBColor(0x1A, 0x36, 0x5D)  # #1a365d - bgzblau from LaTeX
BODY_COLOR = RGBColor(0x00, 0x00, 0x00) # Black body text
GRAY = RGBColor(0x66, 0x66, 0x66)       # Gray for subtitles
FONT_NAME = "Calibri"


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.5):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


def add_horizontal_rule(doc, color="00796b", thickness=1):
    """Add a horizontal rule using a paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{thickness * 8}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def configure_styles(doc):
    """Configure all document styles."""
    styles = doc.styles

    # --- Normal style (body text) ---
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(12)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.space_before = Pt(0)
    # Set east-asia and complex-script font
    rPr = normal._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)

    # --- Heading 1 ---
    h1 = styles["Heading 1"]
    h1.font.name = FONT_NAME
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = TEAL
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h1.paragraph_format.line_spacing = 1.5
    h1.paragraph_format.page_break_before = True
    # Remove underline if any
    h1.font.underline = False

    # --- Heading 2 ---
    h2 = styles["Heading 2"]
    h2.font.name = FONT_NAME
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = TEAL
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h2.paragraph_format.line_spacing = 1.5
    h2.font.underline = False

    # --- Heading 3 ---
    h3 = styles["Heading 3"]
    h3.font.name = FONT_NAME
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = TEAL
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h3.paragraph_format.line_spacing = 1.5
    h3.font.underline = False

    # --- Footnote Text ---
    if "Footnote Text" in styles:
        fn = styles["Footnote Text"]
        fn.font.name = FONT_NAME
        fn.font.size = Pt(10)
        fn.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fn.paragraph_format.space_after = Pt(0)

    # --- TOC styles ---
    for i in range(1, 4):
        toc_name = f"TOC {i}" if f"TOC {i}" in [s.name for s in styles] else None
        if toc_name:
            toc_style = styles[toc_name]
            toc_style.font.name = FONT_NAME
            toc_style.font.size = Pt(12)
            toc_style.font.color.rgb = BODY_COLOR

    # --- Caption style ---
    if "Caption" in styles:
        cap = styles["Caption"]
        cap.font.name = FONT_NAME
        cap.font.size = Pt(10)
        cap.font.italic = True
        cap.font.color.rgb = GRAY

    # --- Create a "Subtitle" style for the title page ---
    if "Subtitle" in styles:
        sub = styles["Subtitle"]
        sub.font.name = FONT_NAME
        sub.font.size = Pt(14)
        sub.font.color.rgb = GRAY

    # --- Create custom "Placeholder" character style ---
    try:
        placeholder_style = styles.add_style("Placeholder", WD_STYLE_TYPE.CHARACTER)
        placeholder_style.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        placeholder_style.font.italic = True
    except ValueError:
        pass  # Style already exists

    # --- Quote / Block Text style ---
    if "Quote" in styles:
        quote = styles["Quote"]
        quote.font.name = FONT_NAME
        quote.font.size = Pt(11)
        quote.font.italic = True
        quote.paragraph_format.left_indent = Cm(1.5)
        quote.paragraph_format.right_indent = Cm(1.5)
        quote.paragraph_format.space_before = Pt(6)
        quote.paragraph_format.space_after = Pt(6)

    # --- List Bullet ---
    if "List Bullet" in styles:
        lb = styles["List Bullet"]
        lb.font.name = FONT_NAME
        lb.font.size = Pt(12)
        lb.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        lb.paragraph_format.line_spacing = 1.5


def configure_page_setup(section):
    """Configure page margins and size."""
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.5)   # 3cm + 0.5cm binding offset
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def add_roman_page_numbers(section):
    """Add Roman numeral page numbers to the footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.name = FONT_NAME
    p.style.font.size = Pt(10)

    # Add page number field with Roman numerals format
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE  \\* ROMAN </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)


def add_arabic_page_numbers(section):
    """Add Arabic page numbers to the footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    run2.font.name = FONT_NAME
    run2.font.size = Pt(10)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE  \\* ARABIC </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    run3.font.name = FONT_NAME
    run3.font.size = Pt(10)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)


def add_header_with_chapter(section):
    """Add header showing chapter name on left, page number on right."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Placeholder text for chapter heading in header
    run = p.add_run()
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY


def restart_page_numbering(section):
    """Restart page numbering at 1 for a section."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")}/>')
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), "1")


def create_title_page(doc):
    """Create the VWA title page matching the LaTeX template."""
    # Spacer at top
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # School logo placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=10, after=4, line_spacing=1.0)
    run = p.add_run("[Schullogo hier einfuegen]")
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # School name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=8, after=2, line_spacing=1.0)
    run = p.add_run("BG Zehnergasse")
    run.font.name = FONT_NAME
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = TEAL

    # School address
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=8, line_spacing=1.0)
    run = p.add_run("Zehnergasse 15, 2700 Wiener Neustadt")
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    # Horizontal rule
    add_horizontal_rule(doc, color="00796b", thickness=1)

    # Document type - "VORWISSENSCHAFTLICHE ARBEIT"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=16, after=2, line_spacing=1.0)
    run = p.add_run("V O R W I S S E N S C H A F T L I C H E")
    run.font.name = FONT_NAME
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = TEAL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.0)
    run = p.add_run("A R B E I T")
    run.font.name = FONT_NAME
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = TEAL

    # Alternative label hint
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=14, line_spacing=1.0)
    run = p.add_run('(oder: Abschliessende Arbeit / ABA)')
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=10, after=6, line_spacing=1.0)
    run = p.add_run("Titel der Arbeit")
    run.font.name = FONT_NAME
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    # Subtitle placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12, line_spacing=1.0)
    run = p.add_run("[Untertitel, falls vorhanden]")
    run.font.name = FONT_NAME
    run.font.size = Pt(13)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # Fachbereich
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=8, after=16, line_spacing=1.0)
    run = p.add_run("Fachbereich: ")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = BODY_COLOR
    run = p.add_run("[z.B. Geschichte und Politische Bildung]")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.italic = True

    # Author section
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=6, after=2, line_spacing=1.0)
    run = p.add_run("Verfasst von:")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = BODY_COLOR

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=2, line_spacing=1.0)
    run = p.add_run("[Vorname Nachname]")
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12, line_spacing=1.0)
    run = p.add_run("Klasse ")
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = BODY_COLOR
    run = p.add_run("[z.B. 8A]")
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.italic = True

    # Supervisor section
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=10, after=2, line_spacing=1.0)
    run = p.add_run("Betreuungsperson:")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = BODY_COLOR

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=12, line_spacing=1.0)
    run = p.add_run("[Mag. Vorname Nachname]")
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.italic = True

    # Location and year
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=10, after=6, line_spacing=1.0)
    run = p.add_run("Wiener Neustadt, ")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = BODY_COLOR
    run = p.add_run("[20XX/XX]")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.italic = True

    # Bottom horizontal rule
    add_horizontal_rule(doc, color="00796b", thickness=1)


def create_declaration(doc):
    """Create the Selbststaendigkeitserklaerung / Eidesstattliche Erklaerung."""
    # Use Heading 1 style but without page break for this special section
    heading = doc.add_heading("Eidesstattliche Erklaerung", level=1)
    heading.paragraph_format.page_break_before = False  # already on new page from break

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=6, line_spacing=1.5)
    run = p.add_run(
        "Ich erklaere, dass ich die vorliegende Vorwissenschaftliche Arbeit "
        "selbststaendig und ohne fremde Hilfe verfasst, andere als die "
        "angegebenen Quellen und Hilfsmittel nicht benutzt und die den "
        "benutzten Quellen woertlich oder inhaltlich entnommenen Stellen als "
        "solche kenntlich gemacht habe."
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
    run = p.add_run(
        'Die Arbeit wurde bisher in gleicher oder aehnlicher Form keiner '
        'anderen Pruefungsbehoerde vorgelegt und auch nicht veroeffentlicht.'
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
    run = p.add_run(
        "Die vorliegende Fassung entspricht der eingereichten elektronischen Version."
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)

    # Date and signature
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=24, after=0, line_spacing=1.5)
    run = p.add_run("Wiener Neustadt, am _______________")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)

    # Signature line
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=36, after=0, line_spacing=1.0)
    run = p.add_run("_________________________________")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=0, line_spacing=1.0)
    run = p.add_run("[Vorname Nachname]")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.italic = True


def create_kurzfassung(doc):
    """Create the German abstract (Kurzfassung)."""
    heading = doc.add_heading("Kurzfassung", level=1)
    heading.paragraph_format.page_break_before = False

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Hier die deutschsprachige Kurzfassung einfuegen (1000-1500 Zeichen). "
        "Die Kurzfassung soll folgende Punkte abdecken: Thema und Forschungsfrage, "
        "Methodik, zentrale Ergebnisse und Schlussfolgerung.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=6, line_spacing=1.5)
    run = p.add_run("Schlagwoerter: ")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.bold = True
    run = p.add_run("[Schlagwort 1, Schlagwort 2, Schlagwort 3, Schlagwort 4]")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.italic = True


def create_abstract(doc):
    """Create the English abstract."""
    heading = doc.add_heading("Abstract", level=1)
    heading.paragraph_format.page_break_before = False

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Insert English abstract here (1000-1500 characters). "
        "The abstract should cover: topic and research question, "
        "methodology, key findings, and conclusion.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=6, line_spacing=1.5)
    run = p.add_run("Keywords: ")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.bold = True
    run = p.add_run("[keyword 1, keyword 2, keyword 3, keyword 4]")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.italic = True


def create_toc_placeholder(doc):
    """Add an auto-generated Table of Contents field."""
    heading = doc.add_heading("Inhaltsverzeichnis", level=1)
    heading.paragraph_format.page_break_before = False

    # Instruction
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=12, line_spacing=1.0)
    run = p.add_run(
        "[Das Inhaltsverzeichnis wird automatisch generiert. "
        "Rechtsklick auf das Feld unten > \"Felder aktualisieren\" "
        "oder Ctrl+A, dann F9 druecken.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # Insert TOC field
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        f' TOC \\o "1-3" \\h \\z \\u '
        f'</w:instrText>'
    )
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar_separate)

    run4 = p.add_run("[Inhaltsverzeichnis - Rechtsklick > Felder aktualisieren]")
    run4.font.name = FONT_NAME
    run4.font.size = Pt(12)
    run4.font.color.rgb = GRAY

    run5 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar2)


def create_section_break(doc, start_type="new_page"):
    """Add a section break for page numbering changes."""
    from docx.enum.section import WD_SECTION_START
    if start_type == "new_page":
        new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    else:
        new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    return new_section


def create_main_body(doc):
    """Create the main body sections with example content."""

    # ---- Chapter 1: Einleitung ----
    doc.add_heading("Einleitung", level=1)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Hinfuehrung zum Thema: Warum ist das Thema relevant? Welche aktuelle "
        "Bedeutung hat es? Hier wird das Interesse der Leserin / des Lesers geweckt "
        "und der Kontext der Arbeit hergestellt.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # 1.1 Forschungsfrage
    doc.add_heading("Forschungsfrage", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "Die zentrale Forschungsfrage dieser Vorwissenschaftlichen Arbeit lautet:"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)

    p = doc.add_paragraph(style="Quote")
    run = p.add_run("[Hier die Forschungsfrage als vollstaendigen Fragesatz formulieren]")
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # 1.2 Methodik
    doc.add_heading("Methodik", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Beschreibung der verwendeten Methodik: z.B. Literaturanalyse, "
        "qualitative Interviews, vergleichende Analyse, empirische Erhebung.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # 1.3 Aufbau der Arbeit
    doc.add_heading("Aufbau der Arbeit", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Die vorliegende Arbeit gliedert sich in X Kapitel. "
        "Kurze Beschreibung des Aufbaus und der einzelnen Kapitel.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # ---- Chapter 2: Theoretische Grundlagen ----
    doc.add_heading("Theoretische Grundlagen", level=1)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Einfuehrung in den theoretischen Rahmen der Arbeit. "
        "Welche Begriffe und Konzepte sind zentral?]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # 2.1 Begriffsbestimmung
    doc.add_heading("Begriffsbestimmung", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Definition und Abgrenzung der wichtigsten Begriffe, "
        "die in der Arbeit verwendet werden.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # 2.2 Stand der Forschung
    doc.add_heading("Stand der Forschung", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Darstellung des aktuellen Forschungsstandes zum Thema. "
        "Welche Arbeiten und Erkenntnisse gibt es bereits?]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # 2.3 Theoretischer Rahmen
    doc.add_heading("Theoretischer Rahmen", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Vorstellung der theoretischen Grundlage, auf der die Arbeit aufbaut. "
        "Welche Modelle oder Theorien werden herangezogen?]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # ---- Chapter 3: Analyse ----
    doc.add_heading("Analyse", level=1)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Einfuehrung in das Analysekapitel. Was wird untersucht und wie?]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # 3.1 Materialauswahl und Vorgehensweise
    doc.add_heading("Materialauswahl und Vorgehensweise", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Beschreibung der verwendeten Materialien (Quellen, Daten, Texte) "
        "und der konkreten Vorgehensweise bei der Analyse.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # 3.2 Durchfuehrung der Analyse
    doc.add_heading("Durchfuehrung der Analyse", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Darstellung der eigentlichen Analyse. Hier werden Argumente "
        "entwickelt, Quellen interpretiert oder Daten ausgewertet.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # 3.3 Ergebnisse
    doc.add_heading("Ergebnisse", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Systematische Praesentation der Ergebnisse. Jedes Teilergebnis "
        "wird klar benannt und mit Belegen untermauert.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # 3.4 Interpretation und Diskussion
    doc.add_heading("Interpretation und Diskussion", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Einordnung der Ergebnisse in den theoretischen Kontext. "
        "Was bedeuten die Ergebnisse? Kritische Reflexion.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # ---- Chapter 4: Fazit ----
    doc.add_heading("Fazit", level=1)

    # 4.1 Zusammenfassung
    doc.add_heading("Zusammenfassung", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Kompakte Zusammenfassung der wichtigsten Inhalte und Erkenntnisse "
        "der Arbeit. Was wurde untersucht? Welche Methode wurde angewendet? "
        "Was sind die zentralen Ergebnisse?]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # 4.2 Beantwortung der Forschungsfrage
    doc.add_heading("Beantwortung der Forschungsfrage", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Explizite und begruendete Beantwortung der Forschungsfrage "
        "auf Grundlage der erarbeiteten Ergebnisse.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # 4.3 Ausblick
    doc.add_heading("Ausblick", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Welche weitergehenden Fragestellungen ergeben sich aus der Arbeit? "
        "Wo besteht weiterer Forschungsbedarf?]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True


def create_bibliography(doc):
    """Create the bibliography section."""
    doc.add_heading("Literaturverzeichnis", level=1)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=12, line_spacing=1.5)
    run = p.add_run(
        "[Alle verwendeten Quellen hier im APA-Stil (7. Auflage) anfuehren. "
        "Beispiele:]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # Example entries
    examples = [
        "Nachname, V. (2023). Titel des Buches (2. Aufl.). Verlag.",
        "Nachname, V. & Nachname, V. (2022). Titel des Artikels. "
        "Name der Zeitschrift, 15(3), 45-67. https://doi.org/10.xxxx",
        "Nachname, V. (2024, 15. Maerz). Titel des Webartikels. "
        "Name der Website. https://www.beispiel.at/artikel",
    ]
    for ex in examples:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)  # hanging indent
        set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
        run = p.add_run(ex)
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.font.color.rgb = GRAY
        run.font.italic = True


def create_appendix(doc):
    """Create the appendix section."""
    doc.add_heading("Anhang", level=1)

    # A.1 Abbildungsverzeichnis
    doc.add_heading("Abbildungsverzeichnis", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Das Abbildungsverzeichnis wird automatisch generiert, wenn "
        "Beschriftungen (Einfuegen > Beschriftung) fuer Abbildungen "
        "verwendet werden. Einfuegen ueber: Verweise > "
        "Abbildungsverzeichnis einfuegen.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # A.2 Tabellenverzeichnis
    doc.add_heading("Tabellenverzeichnis", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Das Tabellenverzeichnis wird automatisch generiert, wenn "
        "Beschriftungen fuer Tabellen verwendet werden.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # A.3 AI Declaration
    doc.add_heading("Erklaerung zum Einsatz von KI-Werkzeugen", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
    run = p.add_run(
        "Gemaess den Richtlinien des BMBWF muss die Verwendung von "
        "KI-gestuetzten Werkzeugen transparent dokumentiert werden."
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
    run = p.add_run(
        "Im Rahmen dieser Vorwissenschaftlichen Arbeit wurden folgende "
        "KI-gestuetzte Werkzeuge eingesetzt:"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)

    # Bullet list for AI tools
    tools_list = [
        ("[Name des Werkzeugs, z.B. ChatGPT 4o]: "
         "[Beschreibung des Einsatzzwecks, z.B. \"Wurde zur Ideenfindung "
         "und Strukturierung von Argumenten verwendet. Alle generierten "
         "Inhalte wurden kritisch geprueft und eigenstaendig ueberarbeitet.\"]"),
        ("[Name des Werkzeugs, z.B. DeepL]: "
         "[Beschreibung des Einsatzzwecks, z.B. \"Wurde zur Uebersetzung "
         "englischsprachiger Fachliteratur als Verstaendnishilfe genutzt.\"]"),
    ]

    for tool in tools_list:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, before=2, after=2, line_spacing=1.5)
        run = p.add_run(tool)
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.font.color.rgb = GRAY
        run.font.italic = True

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Falls keine KI-Werkzeuge eingesetzt wurden: "
        "\"Im Rahmen dieser Vorwissenschaftlichen Arbeit wurden keine "
        "KI-gestuetzten Werkzeuge eingesetzt.\"]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # A.4 Weitere Materialien
    doc.add_heading("Weitere Materialien", level=2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
    run = p.add_run(
        "[Hier koennen weitere Materialien eingefuegt werden, z.B. "
        "Interviewleitfaeden, Transkripte, Fragebogen, statistische "
        "Auswertungen, Abbildungen in hoeherer Aufloesung, etc.]"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.italic = True


def add_formatting_tips(doc):
    """Add a final tips section that students can delete."""
    doc.add_heading("Hinweise zur Verwendung dieser Vorlage", level=1)

    tips = [
        ("Platzhalter ersetzen", (
            "Alle rot-kursiven Textstellen in eckigen Klammern [so wie dieser] "
            "muessen durch eigene Inhalte ersetzt werden."
        )),
        ("Inhaltsverzeichnis aktualisieren", (
            "Nach dem Schreiben: Rechtsklick auf das Inhaltsverzeichnis > "
            "\"Felder aktualisieren\" > \"Gesamtes Verzeichnis aktualisieren\". "
            "Oder: Ctrl+A (alles markieren), dann F9."
        )),
        ("Seitenzahlen", (
            "Der vordere Teil (Erklaerung, Kurzfassung, Inhaltsverzeichnis) "
            "verwendet roemische Ziffern (I, II, III...). Ab der Einleitung "
            "beginnt die arabische Nummerierung (1, 2, 3...)."
        )),
        ("Zitieren im APA-Stil", (
            "Verwenden Sie den APA-Zitierstil (7. Auflage). "
            "Indirekte Zitate: (Nachname, Jahr). "
            "Direkte Zitate: (Nachname, Jahr, S. XX)."
        )),
        ("Abbildungen und Tabellen", (
            "Verwenden Sie Einfuegen > Beschriftung, um Abbildungen und "
            "Tabellen automatisch zu nummerieren. Dies ermoeglicht die "
            "automatische Erstellung eines Abbildungsverzeichnisses."
        )),
        ("Diese Seite loeschen", (
            "Dieses Kapitel \"Hinweise zur Verwendung\" kann nach dem Lesen "
            "vollstaendig geloescht werden."
        )),
    ]

    for title, text in tips:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=8, line_spacing=1.5)
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(12)


def main():
    """Main function to create the VWA Word template."""
    doc = Document()

    # ============================================================
    # Configure styles
    # ============================================================
    configure_styles(doc)

    # ============================================================
    # Section 1: Title page (no page numbers)
    # ============================================================
    section1 = doc.sections[0]
    configure_page_setup(section1)
    # Hide header/footer on title page
    section1.different_first_page_header_footer = False
    footer = section1.footer
    footer.is_linked_to_previous = False
    # Clear footer
    for p in footer.paragraphs:
        p.clear()
    header = section1.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    # Create title page content
    create_title_page(doc)

    # ============================================================
    # Section 2: Front matter with Roman numerals
    # ============================================================
    from docx.enum.section import WD_SECTION_START
    section2 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_page_setup(section2)
    restart_page_numbering(section2)
    add_roman_page_numbers(section2)

    # Declaration
    create_declaration(doc)
    add_page_break(doc)

    # Kurzfassung
    create_kurzfassung(doc)
    add_page_break(doc)

    # Abstract
    create_abstract(doc)
    add_page_break(doc)

    # Table of Contents
    create_toc_placeholder(doc)

    # ============================================================
    # Section 3: Main body with Arabic numerals (restart at 1)
    # ============================================================
    section3 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_page_setup(section3)
    restart_page_numbering(section3)
    add_arabic_page_numbers(section3)

    # Main body content
    create_main_body(doc)

    # Bibliography
    create_bibliography(doc)

    # Appendix
    create_appendix(doc)

    # Usage tips (deletable)
    add_formatting_tips(doc)

    # ============================================================
    # Save the document
    # ============================================================
    output_dir = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "public", "templates"
    )
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "bgz-vwa-word.docx")
    doc.save(output_path)
    print(f"Template saved to: {output_path}")
    print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")


if __name__ == "__main__":
    main()
